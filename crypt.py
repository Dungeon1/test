from tkinter import *
from tkinter import filedialog
import tkinter.scrolledtext as st
import docx2txt
import os
from docx import Document
rm = """1.    Заполните поля, разделяя различные открытые тексты и ключи пустой строкой.

2.    Шифрование осуществляется поблочно, с длиной блока равной длине ключа.
     Если длина ключа больше длины открытого текста (блока), то текст остаётся без изменений.

3.    Кнопка «Выбрать файл» автоматически переносит данные из .docx и .txt файлов в поля программы. 

4.    Если включена галочка «Записать в cipher.docx», то программа создаст указанный документ в директории с программой
    и будет записывать в него выходные данные каждый раз после нажатия кнопки «Зашифровать»."""

def fileopen():
    file = filedialog.askopenfilename(initialdir="/Рабочий стол",title="Select file",filetypes=(("Text document","*.txt *.docx"),("all files","*.*")))
    if file.endswith(".docx"):
        text = docx2txt.process(file)
    else:
        text = open(file,"r")
        text = text.read()
    return text

#Засовывает тест из документа в поля
def get_data():
    text = fileopen()
    plannertexts = []
    keys = []
    for line in text.split("\n\n"):
        if line != "":
            if line[0].isdigit():
                keys.append(line)
            else:
                plannertexts.append(line)
    plannertexts = reversed(plannertexts)
    keys = reversed(keys)
    t1.delete(1.0, END)
    t2.delete(1.0, END)
    for line in plannertexts:                                 
        t1.insert(1.0, line + '\n\n')
    for line in keys:
        t2.insert(1.0, line + '\n\n')
def readme():
    wind = Toplevel(root)
    wind.title('Инструкция')
    display = Label(wind,text=rm,font ='Arial 13',justify =LEFT)
    display.pack()

#Шифрует одну пару текст-ключ
def crypt_one(plannertext, key):
    rule = list(map(int, key.split(' ')))
    rule = [i-1 for i in rule]
    blocks = []
    for x in range(0, len(plannertext), len(rule)):
        blocks.append(plannertext[x:x+len(rule)])
    ciphertext = ""
    for block in blocks:
        try:
            enc_block = [block[i] for i in rule]
            ciphertext += "".join(enc_block)
        except IndexError:
            ciphertext += "".join(block)       
    t3.insert(1.0, ciphertext + '\n\n')
    if cv1.get() == 1:
        exist = os.path.isfile('cipher.docx')
        if exist:
            doc = Document('cipher.docx')
            doc.add_paragraph('Исходный текст:\n'+plannertext+'\nШифртекст:\n'+ciphertext+'\nКлюч: '+key)
            doc.save('cipher.docx')
        else:
            doc = Document()
            doc.add_paragraph('Исходный текст:\n'+plannertext+'\nШифртекст:\n'+ciphertext+'\nКлюч: '+key)
            doc.save('cipher.docx') 

#Шифрует все данные из полей 
def crypt():
    t3.delete(1.0, END)
    plannertexts = []
    keys = []
    for line in t1.get(1.0,'end-1c').split("\n\n"):
        if line != "":
            plannertexts.append(line)
    for line in t2.get(1.0,'end-1c').split("\n\n"):
        if line != "":
            keys.append(line)
    plannertexts = reversed(plannertexts)
    keys = reversed(keys)            
    cyphertexts = list(map(crypt_one, plannertexts, keys))

        
    
#Для контекстного меню ПКМ        
def rClicker(e):
    try:
        def rClick_Copy(e, apnd=0):
            e.widget.event_generate('<Control-c>')

        def rClick_Cut(e):
            e.widget.event_generate('<Control-x>')

        def rClick_Paste(e):
            e.widget.event_generate('<Control-v>')

        e.widget.focus()

        nclst=[
               (' Cut', lambda e=e: rClick_Cut(e)),
               (' Copy', lambda e=e: rClick_Copy(e)),
               (' Paste', lambda e=e: rClick_Paste(e)),
               ]

        rmenu = Menu(None, tearoff=0, takefocus=0)

        for (txt, cmd) in nclst:
            rmenu.add_command(label=txt, command=cmd)

        rmenu.tk_popup(e.x_root+40, e.y_root+10,entry="0")

    except TclError:
        pass

    return "break"


def rClickbinder(r):
    try:
        for b in [ 'Text', 'Entry', 'Listbox', 'Label']: 
            r.bind_class(b, sequence='<Button-3>',
                         func=rClicker, add='')
    except TclError:
        pass

    
root = Tk()
root.title('Шифр перестановки')
#Запускает окно по центру экрана
windowWidth = root.winfo_reqwidth()
windowHeight = root.winfo_reqheight()
positionRight = int(root.winfo_screenwidth()/2 - windowWidth/2)
positionDown = int(root.winfo_screenheight()/2 - windowHeight/2)
root.geometry("+{}+{}".format(positionRight, positionDown))

#Переменные
cv1 = IntVar()

#фрейм с кнопками
f_top = Frame()
b1 = Button(f_top, text="Выбрать файл", width=20, height=3, command=get_data).pack(side=LEFT)
b2 = Button(f_top, text="Зашифровать", width=20, height=3, command=crypt).pack(side=RIGHT)
b3 = Button(f_top,text="Инструкция", width=15, height =2, command=readme).pack()
c1 = Checkbutton(f_top,text="Записать в cipher.docx", variable=cv1, onvalue=1, offvalue=0).pack(side=RIGHT)
f_top.pack()

#фрейм с текстовыми полями
f_bot = Frame()
L1 = Label(f_bot, text="Введите открытый текст").pack(pady=10)
t1 = st.ScrolledText(f_bot, width=50,height=5, font='Arial 13')
t1. insert(1.0, "abcdefg\n\n" + "абвгдеёжз\n\n" + "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.")
t1.pack()

L2 = Label(f_bot, text="Введите правило перестановки").pack(pady=10)
t2 = st.ScrolledText(f_bot, width=50, height=5, font='Arial 13')
t2.insert(1.0, "7 6 5 4 3 2 1\n\n" + "2 3 4 5 6 7 8 9 1\n\n" + "5 2 3 1 4")
t2.pack()

L3 = Label(f_bot, text="Зашифрованное сообщение").pack(pady=10)
t3 = st.ScrolledText(f_bot, width=50, height=5, font='Arial 13')
t3.pack()
f_bot.pack()

rClickbinder(root)
root.mainloop()
