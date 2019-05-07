from tkinter import *
from tkinter import filedialog
import tkinter.scrolledtext as st
import docx2txt
import os
from docx import Document
import random

symbolsAlpha = ['а','б','в','г', 'д', 'е', 'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п',
                'р', 'с', 'т', 'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ', 'ы', 'ь','э', 'ю','я']

symbolsBetta = ['А','Б','В','Г', 'Д', 'Е', 'Ж', 'З', 'И', 'Й', 'К', 'Л', 'М', 'Н', 'О', 'П',
                'Р', 'С', 'Т', 'У', 'Ф', 'Х', 'Ц', 'Ч', 'Ш', 'Щ', 'Ъ', 'Ы', 'Ь','Э', 'Ю','Я']

symbolsRand = ['а','б','в','г', 'д', 'е', 'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п',
                'р', 'с', 'т', 'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ', 'ы', 'ь','э', 'ю','я']

def fileopen():
    file = filedialog.askopenfilename(initialdir="/Рабочий стол",title="Select file",filetypes=(("Text document","*.txt *.docx"),("all files","*.*")))
    if file.endswith(".docx"):
        text = docx2txt.process(file)
    else:
        text = open(file,"r")
        text = text.read()
    return text

def get_data():
    text = fileopen()                                
    t1.insert(1.0, text)

def generator():
    random.shuffle(symbolsRand)
    rand = ''.join(symbolsRand).strip("\n")
    t2.delete(0, END)
    t2.insert(0, rand)

def crypt():
    plannertext = t1.get(1.0, END)
    key = t2.get()
    rule = list(key)
    keys = dict(zip(symbolsAlpha,rule))
    ciphertext=""
    for i in plannertext:
        if i in symbolsAlpha:
            ciphertext += keys[i]
        elif i in symbolsBetta:
            ciphertext += keys[i.lower()].upper()
        else:
            ciphertext += i

    t3.delete(1.0, END)        
    t3.insert(1.0, ciphertext)
    if cv1.get() == 1:
        exist = os.path.isfile('cipher.docx')
        if exist:
            doc = Document('cipher.docx')
            doc.add_paragraph('Исходный текст:\n'+plannertext+'\nШифртекст:\n'+ciphertext+'\nКлюч: '+key)
            doc.save('cipher.docx')
        else:
            doc = Document()
            doc.add_paragraph('Исходный текст:\n'+plannertext+'\nШифртекст:\n'+ciphertext+'\nКлюч: '+key)
            doc.save('cipher.docx')

def clear():
    t1.delete(1.0, END)
    t2.delete(0, END)
    t3.delete(1.0, END)
    
root = Tk()
root.title('Шифр замены')

cv1 = IntVar()

f_top = Frame()
b1 = Button(f_top, text="Выбрать файл", width=15, height=2, command=get_data).pack(side=LEFT)
b2 = Button(f_top, text="Зашифровать", width=15, height=2, command = crypt).pack(side=LEFT)
b4 = Button(f_top, text='Случайный ключ', width=15, height=2, command=generator).pack(side=LEFT)
b5 = Button(f_top, text='Очистить все', width=15, height=2, command=clear).pack(side=RIGHT)
f_top.pack()

c1 = Checkbutton(text="Записать в cipher.docx", variable=cv1, onvalue=1, offvalue=0).pack()

f_text = Frame()
L1 = Label(f_text, text="Введите открытый текст").pack(side=TOP, pady=2)
t1 = st.ScrolledText(f_text, width=60,height=15, font='Arial 13')
t1.pack(side=LEFT)
b5 = Button(f_text, text = 'Очистить', width=8, height=1, command=lambda: t1.delete(1.0, END)).pack(side=RIGHT)
f_text.pack()

f_key = Frame()
L2 = Label(f_key, text="Введите правило замены").pack(side=TOP, pady=2)
t2 = Entry(f_key, width=38, font='Arial 13')
t2.pack(side=LEFT)
b6 = Button(f_key, text = 'Очистить', width=8, height=1, command=lambda: t2.delete(0, END)).pack(side=RIGHT, padx=6)
f_key.pack()

f_cypher = Frame()
L3 = Label(f_cypher, text="Зашифрованное сообщение").pack(side=TOP, pady=2)
t3 = st.ScrolledText(f_cypher, width=60, height=15, font='Arial 13')
t3.pack(side=LEFT)
b7 = Button(f_cypher, text = 'Очистить', width=8, height=1, command=lambda: t3.delete(1.0, END)).pack(side=RIGHT)
f_cypher.pack()

root.mainloop()
