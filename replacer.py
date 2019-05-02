from tkinter import *
from tkinter import filedialog
import tkinter.scrolledtext as st
import docx2txt
import os
from docx import Document
import random
rm = """1. Запишите открытый текст, разделяя различные тексты пустой строкой. Большие буквы переводятся в большие,
символы нерусского алфавита переводятся в себя.

2. Запишите правило замены длиной 32 символа (без 'ё'), отделяя различные правила пустой строкой.

3. Кнопка «Выбрать файл» автоматически переносит данные из .docx и .txt файлов в поля программы. Правило записи данных
   приведено в документе-образце. 

4. Если включена галочка «Записать в cipher.docx», то программа создаст указанный документ в директории с программой
   и будет записывать в него выходные данные каждый раз после нажатия кнопки «Зашифровать»."""
def fileopen():
    file = filedialog.askopenfilename(initialdir="/Рабочий стол",title="Select file",filetypes=(("Text document","*.txt *.docx"),("all files","*.*")))
    if file.endswith(".docx"):
        text = docx2txt.process(file)
    else:
        text = open(file,"r")
        text = text.read()
    return text
symbolsAlpha = ['а','б','в','г', 'д', 'е', 'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п',
                'р', 'с', 'т', 'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ', 'ы', 'ь','э', 'ю','я']
symbolsRand = ['а','б','в','г', 'д', 'е', 'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п',
                'р', 'с', 'т', 'у', 'ф', 'х', 'ц', 'ч', 'ш', 'щ', 'ъ', 'ы', 'ь','э', 'ю','я']
def get_data():
    text = fileopen()
    content = []
    plannertexts = []
    key = []
    content = text.split('\n\n')
    for i in range(0, len(content)):
        if (i%2 == 0):
            plannertexts.append(content[i])
        else:
            key.append(content[i])
            
    plannertexts = reversed(plannertexts)
    key = reversed(key)
    t1.delete(1.0, END)
    t2.delete(1.0, END)
    for line in plannertexts:                                 
        t1.insert(1.0, line + '\n\n')
    for line in key:
        t2.insert(1.0, line + '\n\n')
def readme():
    wind = Toplevel(root)
    wind.title('Инструкция')
    display = Label(wind,text=rm,font ='Arial 13',justify =LEFT)
    display.pack()
def generator():
    random.shuffle(symbolsRand)
    rand = ''.join(symbolsRand).strip("\n")
    t2.insert(INSERT, rand + '\n\n')
def crypt_one(plannertext, key):
    rule = list(key)
    keys = dict(zip(symbolsAlpha,rule))
    ciphertext=""
    for i in plannertext:
        if i.isupper():
            ciphertext+=keys[i.lower()].upper()
        elif i not in symbolsAlpha:
            ciphertext += i
        else:
            ciphertext+=keys[i]

    t3.insert(1.0, ciphertext + '\n\n')
    if cv1.get() == 1:
        exist = os.path.isfile('cipher.docx')
        if exist:
            doc = Document('cipher.docx')
            doc.add_paragraph('Исходный текст:\n'+plannertext+'\nШифртекст:\n'+ciphertext+'\nКлюч: '+key)
            doc.save('cipher.docx')
        else:
            doc = Document()
            doc.add_paragraph('Исходный текст:\n'+plannertext+'\nШифртекст:\n'+ciphertext+'\nКлюч: '+key)
            doc.save('cipher.docx')
def crypt():
    t3.delete(1.0, END)
    plannertexts = []
    keys = []
    for line in t1.get(1.0,'end-1c').split("\n\n"):
        if line != "":
            plannertexts.append(line)
    for line in t2.get(1.0,'end-1c').split("\n\n"):
        if line != "":
            keys.append(line)
    plannertexts = reversed(plannertexts)
    keys = reversed(keys)
    ciphertexts = list(map(crypt_one, plannertexts, keys))

root = Tk()
root.title('Шифр замены')
#Запускает окно по центру экрана

#Переменные
cv1 = IntVar()

#фрейм с кнопками
f_top = Frame()
b1 = Button(f_top, text="Выбрать файл", width=15, height=2, command=get_data).pack(side=LEFT)
b2 = Button(f_top, text="Зашифровать", width=15, height=2, command = crypt).pack(side=LEFT)
b3 = Button(f_top, text="Инструкция", width=15, height =2, command = readme).pack(side=RIGHT)
b4 = Button(f_top, text='Случайный ключ', width=15, height=2, command=generator).pack(side=RIGHT)
f_top.pack()
c1 = Checkbutton(text="Записать в cipher.docx", variable=cv1, onvalue=1, offvalue=0).pack()
#фрейм с текстовыми полями
f_bot = Frame()
L1 = Label(f_bot, text="Введите открытый текст").pack(pady=2)
t1 = st.ScrolledText(f_bot, width=50,height=8, font='Arial 13')
t1. insert(1.0, "абвгдеёжзийклмнопрстуфхцчшщъыьэюя\n\n" + "Это шифр простой замены для символов русского алфавита без буквы йо.\n\n")
t1.pack()

L2 = Label(f_bot, text="Введите правило замены").pack(pady=2)
t2 = st.ScrolledText(f_bot, width=50, height=8, font='Arial 13')
t2.insert(1.0, "яюэьыъщшчцхфутсрпонмлкйизжедгвба\n\n" + "бвгдежзийклмнопрстуфхцчшщъыьэюяа\n\n")
t2.pack()

L3 = Label(f_bot, text="Зашифрованное сообщение").pack(pady=2)
t3 = st.ScrolledText(f_bot, width=50, height=8, font='Arial 13')
t3.pack()
f_bot.pack()

root.mainloop()
